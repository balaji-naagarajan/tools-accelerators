import tkinter as tk
from tkinter import simpledialog, messagebox
import pyautogui
import os
import time
from datetime import datetime
from docx import Document
from docx.shared import Inches
import pygetwindow as gw  # New import for window management

# Global variables to store screenshots and comments
screenshots = []
comments = []

# Function to capture a screenshot
def capture_screenshot():
    # Minimize the tool window
    tool_window = gw.getWindowsWithTitle(root.title())[0]
    tool_window.minimize()
    time.sleep(1)
    
    # Take the screenshot
    screenshot = pyautogui.screenshot()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"screenshot_{timestamp}.png"
    screenshot.save(filename)
    screenshots.append((filename, None, timestamp))  # No comments for this screenshot
    
    # Restore the tool window
    tool_window.restore()
    messagebox.showinfo("Success", f"Screenshot saved as {filename}")

# Function to capture a screenshot with comments
def capture_screenshot_with_comments():
    # Minimize the tool window
    tool_window = gw.getWindowsWithTitle(root.title())[0]
    tool_window.minimize()
    time.sleep(1)
    
    # Take the screenshot
    screenshot = pyautogui.screenshot()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"screenshot_{timestamp}.png"
    screenshot.save(filename)
    
    # Restore the tool window
    tool_window.restore()
    
    # Prompt user for comments
    comment = simpledialog.askstring("Add Comments", "Enter your comments for this screenshot:")
    screenshots.append((filename, comment, timestamp))
    messagebox.showinfo("Success", f"Screenshot with comments saved as {filename}")

# Modify the stop_and_generate_document function to use the user-specified filename
def stop_and_generate_document():
    if not screenshots:
        messagebox.showwarning("No Screenshots", "No screenshots captured to save.")
        return

    # Get the filename from the entry field
    output_filename = filename_entry.get().strip()
    if not output_filename:
        messagebox.showwarning("Invalid Filename", "Please enter a valid filename.")
        return
    output_filename += ".docx"

    # Create a Word document
    doc = Document()
    doc.add_heading(output_filename, level=1)

    for idx, (filename, comment, timestamp) in enumerate(screenshots):
        doc.add_heading(f"Screenshot {idx + 1}", level=2)
        doc.add_paragraph(f"Timestamp: {timestamp}")
        if comment:
            doc.add_paragraph(f"Comments: {comment}")
        doc.add_picture(filename, width=Inches(5))  # Adjust size as needed

    # Save the document
    doc.save(output_filename)

    # Cleanup temporary screenshot files
    for filename, _, _ in screenshots:
        os.remove(filename)

    messagebox.showinfo("Success", f"Document saved as {output_filename}")
    root.destroy()  # Close the application
    
# Create the GUI
root = tk.Tk()
root.title("Screenshot Tool")
root.attributes('-topmost', True)  # Keep the tool window always on top

# Create buttons
btn_capture = tk.Button(root, text="Capture Screenshot", command=capture_screenshot, width=30)
btn_capture_with_comments = tk.Button(root, text="Capture Screenshot with Comments", command=capture_screenshot_with_comments, width=30)
btn_stop = tk.Button(root, text="Stop and Generate Document", command=stop_and_generate_document, width=30)
# Add an entry field for the user to specify the output filename
filename_label = tk.Label(root, text="Enter output filename (without extension):")
filename_label.pack(pady=5)
filename_entry = tk.Entry(root, width=40)
filename_entry.pack(pady=5)

# Place buttons in the window
btn_capture.pack(pady=10)
btn_capture_with_comments.pack(pady=10)
btn_stop.pack(pady=10)

# Run the application
root.mainloop()